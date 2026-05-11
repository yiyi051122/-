# -*- coding: utf-8 -*-
"""
文档解析模块
支持解析txt、pdf、docx等格式文档
"""

import os
from typing import Dict, Optional


class DocumentParser:
    """文档解析器"""
    
    def __init__(self):
        self.supported_extensions = ['.txt', '.pdf', '.docx', '.doc']
    
    def parse(self, file_path: str) -> Dict:
        """
        解析文档
        
        Args:
            file_path: 文档路径
            
        Returns:
            包含text和metadata的字典
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.txt':
            return self._parse_txt(file_path)
        elif ext == '.pdf':
            return self._parse_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")
    
    def _parse_txt(self, file_path: str) -> Dict:
        """解析txt文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        return {
            'text': text,
            'metadata': {
                'file_name': os.path.basename(file_path),
                'file_path': file_path,
                'file_type': 'txt',
                'text_length': len(text)
            }
        }
    
    def _parse_pdf(self, file_path: str) -> Dict:
        """解析PDF文件"""
        try:
            import PyPDF2
            
            text = ""
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            
            return {
                'text': text,
                'metadata': {
                    'file_name': os.path.basename(file_path),
                    'file_path': file_path,
                    'file_type': 'pdf',
                    'text_length': len(text),
                    'num_pages': len(reader.pages)
                }
            }
        except ImportError:
            raise ImportError("请安装PyPDF2: pip install PyPDF2")
    
    def _parse_docx(self, file_path: str) -> Dict:
        """解析Word文档"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            return {
                'text': text,
                'metadata': {
                    'file_name': os.path.basename(file_path),
                    'file_path': file_path,
                    'file_type': 'docx',
                    'text_length': len(text),
                    'num_paragraphs': len(doc.paragraphs)
                }
            }
        except ImportError:
            raise ImportError("请安装python-docx: pip install python-docx")


def parse_document(file_path: str) -> Dict:
    """便捷函数：解析文档"""
    parser = DocumentParser()
    return parser.parse(file_path)


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        result = parse_document(sys.argv[1])
        print(f"文件: {result['metadata']['file_name']}")
        print(f"类型: {result['metadata']['file_type']}")
        print(f"长度: {result['metadata']['text_length']}")
        print(f"内容预览: {result['text'][:200]}...")
